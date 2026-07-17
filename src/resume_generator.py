"""
JobRadar — ATS Resume Generator (Phase 4)

Generates tailored ATS resumes for top-scoring jobs using python-docx.
Only triggers when:
  - overall_score >= 7.5
  - experience_gate != "exclude"
  - daily cap of 5 resumes not reached
  - resume_base.json has been fully populated (sentinel check)
  - budget guard allows the Claude call

Anti-fabrication rule (non-negotiable, §13.1 of brief):
  Every generated resume traces back to resume_base.json line-for-line for
  facts. Only phrasing, ordering, and emphasis may change. No new employers,
  projects, skills, or metrics may be invented.
"""
import json
import logging
import os
import re
from datetime import date
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

logger = logging.getLogger(__name__)

RESUME_BASE_PATH = Path(__file__).parent.parent / "resume_base.json"
DAILY_CAP = 5   # Hard cap — must match config.yaml


# ─────────────────────────────────────────────────────────────────────────────
# Validation
# ─────────────────────────────────────────────────────────────────────────────

def _load_and_validate_resume_base() -> dict:
    """
    Loads resume_base.json and validates that it has been properly populated.
    Raises ValueError if placeholder values are detected.
    """
    if not RESUME_BASE_PATH.exists():
        raise FileNotFoundError(f"resume_base.json not found at {RESUME_BASE_PATH}")

    with open(RESUME_BASE_PATH, encoding="utf-8") as f:
        base = json.load(f)

    sentinel = base.get("_validation_sentinel", "")
    if sentinel != "POPULATED":
        raise ValueError(
            "resume_base.json has not been fully populated. "
            "Set _validation_sentinel to 'POPULATED' after replacing all placeholders."
        )

    # Check for remaining placeholder strings
    base_str = json.dumps(base)
    placeholders = ["YOUR_", "YYYY-MM", "Placeholder —", "YEAR", "YOUR_DEGREE"]
    found = [p for p in placeholders if p in base_str]
    if found:
        raise ValueError(
            f"resume_base.json still contains placeholder values: {found}. "
            "Fill in all real data before enabling resume generation."
        )

    return base


# ─────────────────────────────────────────────────────────────────────────────
# Claude-driven tailoring
# ─────────────────────────────────────────────────────────────────────────────

TAILOR_SYSTEM_PROMPT = """You are an expert ATS resume writer helping Nagesh Khichade tailor his resume for a specific job.

HARD RULES (non-negotiable):
1. You MUST NOT invent any employer, project, skill, tool, certification, or metric not in the provided resume_base.
2. You may ONLY reorder bullets, adjust emphasis, and rephrase existing content to mirror JD terminology.
3. If the JD mentions a skill Nagesh genuinely has (in resume_base.skills), include it prominently.
4. If the JD mentions a skill NOT in resume_base, do NOT claim it — you may mention coursework-adjacent exposure only if it's explicitly in resume_base.skills.coursework_adjacent.
5. The summary should be rewritten to lead with the most relevant skills for THIS job.
6. Experience bullets should reorder/rephrase to emphasise overlap with the JD — never add new facts.

Output: Return a JSON object with these keys:
- "summary": tailored professional summary (2-3 sentences)
- "experience": list of {employer, role, bullets} — same employers from input, reordered bullets
- "projects": list of {name, bullets} — reordered/rephrased, only projects from input
- "skills_to_highlight": list of skills from resume_base that are most relevant to this JD (max 12)
- "cover_note": one sentence suitable for an email cover note (optional, can be empty string)"""


def _tailor_resume_with_claude(
    job: dict,
    resume_base: dict,
    client: anthropic.Anthropic,
    budget_guard,
    config: dict,
) -> dict:
    """Uses Claude Haiku to tailor resume content for a specific job."""
    model = config.get("llm", {}).get("resume_model", "claude-haiku-4-5")
    max_jd_chars = config.get("llm", {}).get("max_jd_chars", 4000)

    est_cost = 0.03  # ~30k tokens max for resume gen
    budget_guard.check_and_debit("claude_resume", est_cost)

    prompt = f"""Job to tailor resume for:
Title: {job.get('title', '')}
Company: {job.get('company', '')}
Location: {job.get('location', '')}
Job Description:
{(job.get('description_text', '') or '')[:max_jd_chars]}

Resume base (source of truth — do not add facts not in here):
{json.dumps(resume_base, indent=2, ensure_ascii=False)[:6000]}

Tailor the resume for this specific job following all HARD RULES."""

    response = client.messages.create(
        model=model,
        max_tokens=2048,
        system=TAILOR_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = re.sub(r"```[a-z]*\n?", "", raw).strip().rstrip("```").strip()

    return json.loads(raw)


# ─────────────────────────────────────────────────────────────────────────────
# python-docx resume builder
# ─────────────────────────────────────────────────────────────────────────────

def _build_docx(tailored: dict, resume_base: dict, job: dict, output_path: Path) -> Path:
    """Builds a clean ATS-friendly .docx from tailored content."""
    doc = Document()

    personal = resume_base.get("personal", {})

    # ── Style helpers ─────────────────────────────────────────────────────────
    def add_heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
        return p

    def set_font(paragraph, size_pt: int = 11):
        for run in paragraph.runs:
            run.font.size = Pt(size_pt)

    # ── Page margins (narrow for ATS) ────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # ── Header ────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(personal.get("name", "Nagesh Khichade"))
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_parts = [
        personal.get("email", ""),
        personal.get("phone", ""),
        personal.get("location", ""),
        personal.get("linkedin", ""),
        personal.get("github", ""),
    ]
    contact_str = "  |  ".join(p for p in contact_parts if p)
    contact_para = doc.add_paragraph(contact_str)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ── Summary ───────────────────────────────────────────────────────────────
    add_heading("Professional Summary", level=2)
    doc.add_paragraph(tailored.get("summary", resume_base.get("summary", "")))

    # ── Skills ────────────────────────────────────────────────────────────────
    add_heading("Skills", level=2)
    highlighted = tailored.get("skills_to_highlight", [])
    all_skills = resume_base.get("skills", {})
    all_skill_flat = (
        all_skills.get("languages", [])
        + all_skills.get("ml_ai", [])
        + all_skills.get("data_engineering", [])
        + all_skills.get("analytics_viz", [])
        + all_skills.get("databases", [])
        + all_skills.get("tools", [])
        + all_skills.get("cloud", [])
    )
    display_skills = highlighted if highlighted else all_skill_flat[:15]
    doc.add_paragraph(", ".join(display_skills))

    # ── Experience ────────────────────────────────────────────────────────────
    add_heading("Experience", level=2)
    tailored_exp = tailored.get("experience", resume_base.get("experience", []))
    for exp in tailored_exp:
        employer = exp.get("employer", "")
        role = exp.get("role", "")
        p = doc.add_paragraph()
        p.add_run(f"{role} — {employer}").bold = True
        for bullet in exp.get("bullets", []):
            add_bullet(bullet)

    # ── Projects ─────────────────────────────────────────────────────────────
    add_heading("Projects", level=2)
    tailored_projects = tailored.get("projects", [])
    base_projects = {p["name"]: p for p in resume_base.get("projects", [])}

    for proj in tailored_projects:
        pname = proj.get("name", "")
        base_proj = base_projects.get(pname, {})
        tech = base_proj.get("tech", [])
        p = doc.add_paragraph()
        p.add_run(pname).bold = True
        if tech:
            p.add_run(f"  [{', '.join(tech[:6])}]").italic = True
        for bullet in proj.get("bullets", base_proj.get("bullets", [])):
            add_bullet(bullet)

    # ── Education ────────────────────────────────────────────────────────────
    add_heading("Education", level=2)
    for edu in resume_base.get("education", []):
        p = doc.add_paragraph()
        p.add_run(f"{edu.get('degree', '')} — {edu.get('institution', '')}").bold = True
        year = edu.get("year", "")
        if year:
            p.add_run(f" ({year})")
        coursework = edu.get("relevant_coursework", [])
        if coursework:
            doc.add_paragraph(f"Coursework: {', '.join(coursework)}")

    # ── Certifications ────────────────────────────────────────────────────────
    certs = resume_base.get("certifications", [])
    if certs:
        add_heading("Certifications", level=2)
        for cert in certs:
            doc.add_paragraph(cert, style="List Bullet")

    doc.save(str(output_path))
    logger.info(f"Resume saved: {output_path}")
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Google Drive upload
# ─────────────────────────────────────────────────────────────────────────────

def _upload_to_drive(file_path: Path, folder_id: str) -> Optional[str]:
    """Uploads a file to Google Drive and returns the shareable link."""
    try:
        sa_json = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
        if not sa_json:
            raise EnvironmentError("GOOGLE_SERVICE_ACCOUNT_JSON not set")

        creds_info = json.loads(sa_json)
        creds = Credentials.from_service_account_info(
            creds_info,
            scopes=["https://www.googleapis.com/auth/drive"],
        )
        service = build("drive", "v3", credentials=creds)

        file_metadata = {
            "name": file_path.name,
            "parents": [folder_id],
        }
        media = MediaFileUpload(
            str(file_path),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        uploaded = service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id,webViewLink",
        ).execute()

        link = uploaded.get("webViewLink", "")
        logger.info(f"Uploaded to Drive: {link}")
        return link
    except Exception as e:
        logger.error(f"Drive upload failed: {e}")
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def generate_resumes(
    jobs: list[dict],
    anthropic_client: anthropic.Anthropic,
    budget_guard,
    config: dict,
) -> list[dict]:
    """
    Generates tailored ATS resumes for top-scoring jobs.
    Modifies jobs in-place with 'resume_link' field.
    Hard cap: 5 resumes per day.
    """
    min_score = config.get("scoring", {}).get("resume_min_score", 7.5)
    max_per_day = config.get("scoring", {}).get("resume_max_per_day", DAILY_CAP)
    drive_folder_id = os.environ.get("GOOGLE_DRIVE_FOLDER_ID", "")

    # Validate resume_base.json before spending any budget
    try:
        resume_base = _load_and_validate_resume_base()
    except (FileNotFoundError, ValueError) as e:
        logger.warning(f"Resume generation skipped: {e}")
        for job in jobs:
            job["resume_link"] = ""
        return jobs

    # Filter eligible jobs
    eligible = [
        j for j in jobs
        if j.get("overall_score", 0) >= min_score
        and j.get("experience_gate_label", "") != "exclude"
    ]
    # Sort by score descending, take top max_per_day
    eligible.sort(key=lambda j: j.get("overall_score", 0), reverse=True)
    eligible = eligible[:max_per_day]
    logger.info(f"Resume gen: {len(eligible)} eligible (score >= {min_score}, cap {max_per_day})")

    today_str = date.today().strftime("%Y-%m-%d")
    output_dir = Path("resumes")
    output_dir.mkdir(exist_ok=True)

    for job in eligible:
        try:
            tailored = _tailor_resume_with_claude(job, resume_base, anthropic_client, budget_guard, config)

            company_slug = re.sub(r"[^A-Za-z0-9]", "_", job.get("company", "Unknown"))[:20]
            title_slug = re.sub(r"[^A-Za-z0-9]", "_", job.get("title", "Role"))[:20]
            filename = f"Nagesh_Khichade_{company_slug}_{title_slug}_{today_str}.docx"
            output_path = output_dir / filename

            _build_docx(tailored, resume_base, job, output_path)

            if drive_folder_id:
                link = _upload_to_drive(output_path, drive_folder_id)
                job["resume_link"] = link or ""
            else:
                job["resume_link"] = str(output_path)

            logger.info(f"Resume generated for: {job.get('title')} @ {job.get('company')}")

        except Exception as e:
            logger.error(f"Resume gen failed for job {job.get('job_id')}: {e}")
            job["resume_link"] = ""

    # Jobs that didn't get a resume
    generated_ids = {j["job_id"] for j in eligible}
    for job in jobs:
        if job["job_id"] not in generated_ids:
            job.setdefault("resume_link", "")

    return jobs
